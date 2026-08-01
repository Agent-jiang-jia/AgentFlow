"""Real supported-format parser tests."""

from pathlib import Path

import fitz  # type: ignore[import-untyped]  # PyMuPDF has no typing metadata.
from app.parsers.csv import CsvParser
from app.parsers.docx import DocxParser
from app.parsers.pdf import PdfParser
from app.parsers.text import TextParser
from docx import Document


def test_text_parser_falls_back_to_common_chinese_encoding(tmp_path: Path) -> None:
    """GB18030 text is normalized to bounded UTF-8 Markdown."""
    source = tmp_path / "中文.txt"
    source.write_bytes("第一段\r\n\r\n第二段".encode("gb18030"))
    result = TextParser().parse(source, original_name="中文.txt", max_chars=10_000)
    assert result.status == "success"
    assert result.content == "# 文件: 中文.txt\n\n第一段\n\n第二段\n"


def test_csv_parser_reports_total_and_limits_to_500_rows(tmp_path: Path) -> None:
    """CSV output preserves a header while explicitly marking row truncation."""
    source = tmp_path / "data.csv"
    rows = ["name,value", *(f"row-{index},{index}" for index in range(501))]
    source.write_text("\n".join(rows), encoding="utf-8")
    result = CsvParser().parse(source, original_name="data.csv", max_chars=100_000)
    assert result.content is not None
    assert "总数据行数: 501" in result.content
    assert "是否截断: 是" in result.content
    assert "row-499" in result.content
    assert "row-500" not in result.content


def test_docx_parser_extracts_heading_list_and_table(tmp_path: Path) -> None:
    """DOCX conversion includes the required common document structures."""
    source = tmp_path / "requirements.docx"
    document = Document()
    document.add_heading("需求说明", level=1)
    document.add_paragraph("核心功能")
    document.add_paragraph("安全上传", style="List Bullet")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "模块"
    table.cell(0, 1).text = "风险"
    table.cell(1, 0).text = "文件"
    table.cell(1, 1).text = "路径|穿越"
    document.save(str(source))

    result = DocxParser().parse(source, original_name="requirements.docx", max_chars=10_000)
    assert result.content is not None
    assert "## 需求说明" in result.content
    assert "- 安全上传" in result.content
    assert "| 模块 | 风险 |" in result.content
    assert r"路径\|穿越" in result.content


def test_pdf_parser_extracts_pages_and_identifies_scan(tmp_path: Path) -> None:
    """Text PDFs become page-aware Markdown while blank scans do not create content."""
    text_pdf = tmp_path / "text.pdf"
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "This PDF contains enough selectable text for parsing.")
    document.save(text_pdf)
    document.close()
    parsed = PdfParser().parse(text_pdf, original_name="text.pdf", max_chars=10_000)
    assert parsed.status == "success"
    assert parsed.content is not None
    assert "## 第 1 页" in parsed.content

    scan_pdf = tmp_path / "scan.pdf"
    document = fitz.open()
    document.new_page()
    document.save(scan_pdf)
    document.close()
    scanned = PdfParser().parse(scan_pdf, original_name="scan.pdf", max_chars=10_000)
    assert scanned.status == "unsupported_ocr"
    assert scanned.content is None
    assert scanned.error is not None
