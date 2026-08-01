"""DOCX paragraphs, lists, headings, and table extraction."""

from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject

from app.parsers.base import ParseResult, markdown_cell, truncate_text


class DocxParser:
    """Convert common DOCX document structures into Markdown."""

    def parse(self, path: Path, *, original_name: str, max_chars: int) -> ParseResult:
        """Extract visible paragraphs and tables without interpreting embedded media."""
        document: DocumentObject = Document(str(path))
        blocks = [f"# 文件: {original_name}"]
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = paragraph.style.name.lower() if paragraph.style is not None else ""
            if style_name.startswith("heading"):
                suffix = style_name.removeprefix("heading").strip()
                level = min(6, max(2, int(suffix) + 1 if suffix.isdigit() else 2))
                blocks.append(f"{'#' * level} {text}")
            elif style_name.startswith("list"):
                blocks.append(f"- {text}")
            else:
                blocks.append(text)

        for table in document.tables:
            rows = [[markdown_cell(cell.text.strip()) for cell in row.cells] for row in table.rows]
            if not rows:
                continue
            column_count = max(len(row) for row in rows)
            padded = [row + [""] * (column_count - len(row)) for row in rows]
            header = padded[0]
            blocks.append(
                "\n".join(
                    [
                        f"| {' | '.join(header)} |",
                        f"| {' | '.join(['---'] * column_count)} |",
                        *(f"| {' | '.join(row)} |" for row in padded[1:]),
                    ]
                )
            )

        if len(blocks) == 1:
            raise ValueError("DOCX contains no readable text")
        return ParseResult(
            status="success",
            content=truncate_text("\n\n".join(blocks).strip() + "\n", max_chars),
        )
